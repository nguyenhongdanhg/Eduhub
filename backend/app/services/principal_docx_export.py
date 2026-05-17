from __future__ import annotations

import io
import re
import unicodedata
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.services.principal_document_presets import get_principal_document_preset

CONTENT_TYPE_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _plain(s: Any) -> str:
  return str(s or "").replace("\r\n", "\n").replace("\r", "\n").strip()


def _compact(s: Any) -> str:
  return re.sub(r"\s+", " ", str(s or "")).strip()


def _strip_markdown(s: str) -> str:
  t = str(s or "")
  t = re.sub(r"^#{1,6}\s+", "", t)
  t = re.sub(r"\*\*(.*?)\*\*", r"\1", t)
  t = re.sub(r"__(.*?)__", r"\1", t)
  t = re.sub(r"`([^`]*)`", r"\1", t)
  return t.strip()


def _slug(s: str, fallback: str = "van-ban") -> str:
  t = unicodedata.normalize("NFD", str(s or ""))
  t = "".join(ch for ch in t if unicodedata.category(ch) != "Mn")
  t = re.sub(r"[^A-Za-z0-9]+", "-", t).strip("-").lower()
  return t[:80] or fallback


def _filename(decision_id: int, title: str, prefix: str) -> str:
  p = _slug(prefix or "van-ban")
  t = _slug(title or "du-thao")
  return f"{p}-{int(decision_id or 0)}-{t}.docx"


def _set_cell_text(cell, text: str, *, bold: bool = False, italic: bool = False, size: int = 13, align=None):
  cell.text = ""
  p = cell.paragraphs[0]
  if align is not None:
    p.alignment = align
  r = p.add_run(text)
  r.bold = bool(bold)
  r.italic = bool(italic)
  r.font.name = "Times New Roman"
  r.font.size = Pt(size)


def _set_cell_border_none(cell):
  tc = cell._tc
  tc_pr = tc.get_or_add_tcPr()
  borders = tc_pr.first_child_found_in("w:tcBorders")
  if borders is None:
    borders = OxmlElement("w:tcBorders")
    tc_pr.append(borders)
  for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
    tag = "w:" + edge
    el = borders.find(qn(tag))
    if el is None:
      el = OxmlElement(tag)
      borders.append(el)
    el.set(qn("w:val"), "nil")


def _style_doc(doc: Document):
  sec = doc.sections[0]
  sec.start_type = WD_SECTION_START.NEW_PAGE
  sec.page_width = Cm(21.0)
  sec.page_height = Cm(29.7)
  sec.top_margin = Cm(2.0)
  sec.bottom_margin = Cm(2.0)
  sec.left_margin = Cm(3.0)
  sec.right_margin = Cm(1.5)
  sec.header_distance = Cm(1.0)
  sec.footer_distance = Cm(1.0)

  styles = doc.styles
  normal = styles["Normal"]
  normal.font.name = "Times New Roman"
  normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
  normal.font.size = Pt(13)
  normal.paragraph_format.line_spacing = 1.15
  normal.paragraph_format.space_after = Pt(3)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: int = 13, underline: bool = False):
  run = paragraph.add_run(text)
  run.bold = bold
  run.italic = italic
  run.underline = underline
  run.font.name = "Times New Roman"
  run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
  run.font.size = Pt(size)
  return run


def _add_paragraph(doc: Document, text: str = "", *, align=None, bold: bool = False, italic: bool = False, size: int = 13, space_after: int = 3, first_line: bool = False):
  p = doc.add_paragraph()
  if align is not None:
    p.alignment = align
  if first_line:
    p.paragraph_format.first_line_indent = Cm(1.0)
  p.paragraph_format.space_after = Pt(space_after)
  p.paragraph_format.line_spacing = 1.15
  if text:
    _add_run(p, text, bold=bold, italic=italic, size=size)
  return p


def _add_header_blocks(doc: Document, meta: dict, preset: dict):
  agency = _compact(meta.get("agency_name") or meta.get("co_quan") or "[TÊN CƠ QUAN, ĐƠN VỊ]")
  doc_no = _compact(meta.get("document_no") or meta.get("so_ky_hieu") or "Số: [số]/[ký hiệu]")
  place = _compact(meta.get("place") or meta.get("dia_danh") or "[Địa danh]")
  date_text = _compact(meta.get("date_text") or meta.get("ngay_thang") or datetime.now().strftime("ngày %d tháng %m năm %Y"))

  table = doc.add_table(rows=1, cols=2)
  table.alignment = WD_TABLE_ALIGNMENT.CENTER
  table.columns[0].width = Cm(8.0)
  table.columns[1].width = Cm(9.0)
  for cell in table.rows[0].cells:
    _set_cell_border_none(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

  left, right = table.rows[0].cells
  _set_cell_text(left, agency.upper(), bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
  p = left.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  _add_run(p, doc_no, size=12)

  _set_cell_text(right, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
  p2 = right.add_paragraph()
  p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
  _add_run(p2, "Độc lập - Tự do - Hạnh phúc", bold=True, underline=True, size=12)
  p3 = right.add_paragraph()
  p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
  _add_run(p3, f"{place}, {date_text}", italic=True, size=12)

  _add_paragraph(doc, "", space_after=6)


def _add_title(doc: Document, title: str, preset: dict):
  doc_type = _compact(preset.get("doc_type") or preset.get("label") or "DỰ THẢO VĂN BẢN").upper()
  _add_paragraph(doc, doc_type, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=3)
  title2 = _compact(title)
  if title2 and title2.upper() != doc_type:
    _add_paragraph(doc, title2, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, space_after=8)


def _is_heading(line: str) -> bool:
  t = line.strip()
  if not t:
    return False
  if re.match(r"^#{1,6}\s+", t):
    return True
  if re.match(r"^(I|II|III|IV|V|VI|VII|VIII|IX|X)\.\s+", t, flags=re.I):
    return True
  if re.match(r"^(Điều|Chương|Phần|Mục)\s+\d+", t, flags=re.I):
    return True
  if len(t) <= 80 and t.endswith(":"):
    return True
  return False


def _add_body_from_text(doc: Document, generated_text: str):
  text = _plain(generated_text)
  lines = text.split("\n")
  blank_count = 0
  for raw in lines:
    line = raw.strip()
    if not line:
      blank_count += 1
      if blank_count <= 1:
        _add_paragraph(doc, "", space_after=2)
      continue
    blank_count = 0
    bullet = re.match(r"^[-+*•]\s+(.*)$", line)
    numbered = re.match(r"^(\d+[\.)]|[a-zA-Z][\.)])\s+(.*)$", line)
    if bullet:
      p = _add_paragraph(doc, _strip_markdown(bullet.group(1)), size=13, space_after=2)
      p.paragraph_format.left_indent = Cm(0.8)
      p.paragraph_format.first_line_indent = Cm(-0.3)
      p.runs[0].text = "- " + p.runs[0].text
      continue
    if numbered:
      p = _add_paragraph(doc, _strip_markdown(line), size=13, space_after=2)
      p.paragraph_format.left_indent = Cm(0.7)
      continue
    heading = _is_heading(line)
    cleaned = _strip_markdown(line)
    _add_paragraph(doc, cleaned, bold=heading, size=13 if not heading else 13, space_after=4, first_line=not heading)


def _add_signature(doc: Document, meta: dict):
  signer_title = _compact(meta.get("signer_title") or meta.get("chuc_vu_ky") or "HIỆU TRƯỞNG")
  signer_name = _compact(meta.get("signer_name") or meta.get("nguoi_ky") or "[Họ và tên]")
  recipients = _plain(meta.get("recipients") or meta.get("noi_nhan") or "- Như trên;\n- Lưu: VT.")

  _add_paragraph(doc, "", space_after=4)
  table = doc.add_table(rows=1, cols=2)
  table.columns[0].width = Cm(8.5)
  table.columns[1].width = Cm(8.5)
  for cell in table.rows[0].cells:
    _set_cell_border_none(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

  left, right = table.rows[0].cells
  _set_cell_text(left, "Nơi nhận:", bold=True, italic=True, size=12)
  for ln in recipients.split("\n"):
    t = ln.strip()
    if t:
      p = left.add_paragraph()
      _add_run(p, t, size=11)

  _set_cell_text(right, signer_title.upper(), bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
  p = right.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  _add_run(p, "(Ký, ghi rõ họ tên)", italic=True, size=11)
  for _ in range(4):
    right.add_paragraph()
  p2 = right.add_paragraph()
  p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
  _add_run(p2, signer_name, bold=True, size=13)


def _add_review_notes(doc: Document, generated_text: str):
  if "Rà soát trước khi ban hành" in generated_text:
    return
  _add_paragraph(doc, "", space_after=4)
  _add_paragraph(doc, "Rà soát trước khi ban hành", bold=True, size=12, space_after=2)
  for item in (
    "Kiểm tra lại số/ký hiệu, ngày tháng và thẩm quyền ký.",
    "Bổ sung các thông tin còn để trong ngoặc [cần bổ sung] nếu có.",
    "Đối chiếu căn cứ, đơn vị thực hiện và thời hạn trước khi phát hành.",
  ):
    p = _add_paragraph(doc, item, size=11, italic=True, space_after=1)
    p.paragraph_format.left_indent = Cm(0.8)
    if p.runs:
      p.runs[0].text = "- " + p.runs[0].text


def export_principal_decision_docx(*, title: str, generated_text: str, meta: dict, decision_id: int) -> tuple[bytes, str]:
  text = _plain(generated_text)
  if not text:
    raise RuntimeError("empty_generated_text")
  meta2 = meta if isinstance(meta, dict) else {}
  preset = get_principal_document_preset(meta2.get("preset_id"))
  doc = Document()
  _style_doc(doc)
  _add_header_blocks(doc, meta2, preset)
  _add_title(doc, title or meta2.get("title") or meta2.get("user_request") or "Dự thảo văn bản", preset)
  _add_body_from_text(doc, text)
  _add_signature(doc, meta2)
  _add_review_notes(doc, text)
  buf = io.BytesIO()
  doc.save(buf)
  filename = _filename(int(decision_id or 0), title or meta2.get("title") or "du-thao", str(preset.get("filename_prefix") or "van-ban"))
  return buf.getvalue(), filename
